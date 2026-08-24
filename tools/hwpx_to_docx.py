from __future__ import annotations

import argparse
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from lxml import etree

from tools.docx_footnotes import add_footnotes
from tools.hwpx_docx_format import FormatCatalog
from tools.hwpx_docx_format import direct as _direct
from tools.hwpx_docx_format import first as _first
from tools.hwpx_docx_format import hwp_inches as _hwp_inches
from tools.hwpx_docx_format import local_name as _local
from tools.hwpx_docx_format import number as _number


class Converter:
    def __init__(self, source: Path) -> None:
        self.source = source
        self.document: DocumentObject = Document()
        self.formats: FormatCatalog | None = None
        self.footnotes: list[list[str]] = []
        self.assets: dict[str, Path] = {}
        self._temporary: tempfile.TemporaryDirectory[str] | None = None

    def load(self) -> etree._Element:
        self._temporary = tempfile.TemporaryDirectory(prefix="hwpx-docx-")
        root = Path(self._temporary.name)
        with zipfile.ZipFile(self.source) as archive:
            archive.extractall(root)
        header = etree.parse(root / "Contents/header.xml").getroot()
        section = etree.parse(root / "Contents/section0.xml").getroot()
        self.formats = FormatCatalog(header)
        self.assets = {
            path.stem: path for path in (root / "BinData").glob("*") if path.is_file()
        }
        return section

    def add_footnote_reference(
        self, paragraph: Paragraph, note: etree._Element
    ) -> None:
        note_id = len(self.footnotes) + 1
        paragraphs = [
            "".join(str(text) for text in node.itertext())
            for node in note.iter()
            if _local(node) == "p"
        ]
        self.footnotes.append([text for text in paragraphs if text])
        run = paragraph.add_run()
        properties = run._r.get_or_add_rPr()
        etree.SubElement(properties, qn("w:rStyle"), {qn("w:val"): "FootnoteReference"})
        etree.SubElement(run._r, qn("w:footnoteReference"), {qn("w:id"): str(note_id)})
        suffix = OxmlElement("w:t")
        suffix.text = ")"
        run._r.append(suffix)

    def fill_paragraph(self, paragraph: Paragraph, source: etree._Element) -> None:
        if self.formats is None:
            raise RuntimeError("Document formats were not loaded")
        self.formats.apply_paragraph(paragraph, source.get("paraPrIDRef", ""))
        if source.get("pageBreak") == "1":
            paragraph.paragraph_format.page_break_before = True
        for run_source in _direct(source, "run"):
            style_id = run_source.get("charPrIDRef", "")
            for child in run_source:
                kind = _local(child)
                if kind == "t":
                    run = paragraph.add_run()
                    self.formats.apply_run(run, style_id)
                    if child.text:
                        run.add_text(child.text)
                    for inline in child:
                        if _local(inline) == "lineBreak":
                            run.add_break()
                        elif _local(inline) == "tab":
                            run.add_tab()
                        if inline.tail:
                            run.add_text(inline.tail)
                elif kind == "ctrl":
                    for control in child:
                        if _local(control) == "footNote":
                            self.add_footnote_reference(paragraph, control)
                elif kind == "pic":
                    image = _first(child, "img")
                    size = _first(child, "curSz")
                    image_id = (
                        image.get("binaryItemIDRef", "") if image is not None else ""
                    )
                    if image_id in self.assets:
                        width = _hwp_inches(_number(size, "width"))
                        paragraph.add_run().add_picture(
                            str(self.assets[image_id]), width=width
                        )

    def fill_cell(self, cell: _Cell, source: etree._Element) -> None:
        sublist = _first(source, "subList")
        if sublist is None:
            return
        cell.vertical_alignment = {
            "TOP": WD_CELL_VERTICAL_ALIGNMENT.TOP,
            "CENTER": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            "BOTTOM": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
        }.get(sublist.get("vertAlign", ""), WD_CELL_VERTICAL_ALIGNMENT.TOP)
        if self.formats is None:
            raise RuntimeError("Document formats were not loaded")
        self.formats.apply_cell(cell, source)
        paragraphs = _direct(sublist, "p")
        for index, paragraph_source in enumerate(paragraphs):
            paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
            self.fill_paragraph(paragraph, paragraph_source)

    def add_table(self, source: etree._Element) -> Table:
        rows = _number(source, "rowCnt", 1)
        columns = _number(source, "colCnt", 1)
        table = self.document.add_table(rows=rows, cols=columns)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        if self.formats is None:
            raise RuntimeError("Document formats were not loaded")
        self.formats.apply_table(table, source)
        cell_sources = [node for node in source.iter() if _local(node) == "tc"]
        column_widths = [0] * columns
        for cell_source in cell_sources:
            address = _first(cell_source, "cellAddr")
            span = _first(cell_source, "cellSpan")
            size = _first(cell_source, "cellSz")
            column = _number(address, "colAddr")
            column_span = _number(span, "colSpan", 1)
            if column_span == 1:
                column_widths[column] = max(
                    column_widths[column], _number(size, "width")
                )
        table_size = _first(source, "sz")
        fallback_width = _number(table_size, "width", 36000) // columns
        for index, width in enumerate(column_widths):
            table.columns[index].width = _hwp_inches(width or fallback_width)
        for cell_source in cell_sources:
            address = _first(cell_source, "cellAddr")
            span = _first(cell_source, "cellSpan")
            row = _number(address, "rowAddr")
            column = _number(address, "colAddr")
            row_span = _number(span, "rowSpan", 1)
            column_span = _number(span, "colSpan", 1)
            cell = table.cell(row, column)
            if row_span > 1 or column_span > 1:
                cell = cell.merge(
                    table.cell(row + row_span - 1, column + column_span - 1)
                )
            self.fill_cell(cell, cell_source)
        return table

    def convert(self, target: Path) -> None:
        section = self.load()
        FormatCatalog.configure_page(self.document, section)
        for source in _direct(section, "p"):
            tables = [node for node in source.iter() if _local(node) == "tbl"]
            has_content = any(
                _local(node) in {"t", "footNote", "pic"} for node in source.iter()
            )
            if has_content or not tables:
                paragraph = self.document.add_paragraph()
                self.fill_paragraph(paragraph, source)
            for table_source in tables:
                self.add_table(table_source)
        target.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(target))
        add_footnotes(target, self.footnotes)
        if self._temporary is not None:
            self._temporary.cleanup()


def convert_hwpx(source: Path, target: Path) -> None:
    Converter(source).convert(target)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert HWPX to DOCX locally.")
    parser.add_argument("source", type=Path)
    parser.add_argument("target", type=Path)
    arguments = parser.parse_args()
    convert_hwpx(arguments.source, arguments.target)


if __name__ == "__main__":
    main()
