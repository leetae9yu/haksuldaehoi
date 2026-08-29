from __future__ import annotations

import copy

from docx.table import Table
from lxml import etree

from tools.docx_hwpx_content import direct, fill_paragraph, local_name
from tools.docx_hwpx_styles import CharCatalog


def fill_table(
    target: etree._Element,
    source: Table,
    catalog: CharCatalog,
    notes: dict[str, str],
    note_template: etree._Element,
    note_number: list[int],
) -> None:
    for cell_source in (
        node for node in target.iter() if local_name(node) == "tc"
    ):
        address = next(
            child
            for child in cell_source
            if local_name(child) == "cellAddr"
        )
        row = int(address.get("rowAddr", "0"))
        column = int(address.get("colAddr", "0"))
        source_cell = source.cell(row, column)
        sublist = next(
            child for child in cell_source if local_name(child) == "subList"
        )
        paragraphs = direct(sublist, "p")
        while len(paragraphs) < len(source_cell.paragraphs):
            clone = copy.deepcopy(paragraphs[-1])
            sublist.append(clone)
            paragraphs.append(clone)
        for paragraph in paragraphs[len(source_cell.paragraphs) :]:
            sublist.remove(paragraph)
        paragraphs = paragraphs[: len(source_cell.paragraphs)]
        for target_paragraph, source_paragraph in zip(
            paragraphs,
            source_cell.paragraphs,
            strict=False,
        ):
            fill_paragraph(
                target_paragraph,
                source_paragraph,
                catalog,
                notes,
                note_template,
                note_number,
            )


def append_single_cell_table(
    section: etree._Element,
    table_paragraph: etree._Element,
    table_template: etree._Element,
    paragraph_template: etree._Element,
    table_id: str,
    width: str,
) -> etree._Element:
    paragraph = copy.deepcopy(table_paragraph)
    anchor_runs = direct(paragraph, "run")
    anchor_style = (
        anchor_runs[0].get("charPrIDRef", "0") if anchor_runs else "0"
    )
    for anchor_run in anchor_runs:
        paragraph.remove(anchor_run)
    table = copy.deepcopy(table_template)
    table.set("id", table_id)
    table.set("rowCnt", "1")
    table.set("colCnt", "1")
    size = next(node for node in table if local_name(node) == "sz")
    size.set("width", width)
    rows = [node for node in table if local_name(node) == "tr"]
    body_cells = [
        copy.deepcopy(node)
        for node in table.iter()
        if local_name(node) == "tc"
        and int(
            next(
                child
                for child in node
                if local_name(child) == "cellAddr"
            ).get("rowAddr", "0")
        )
        > 0
    ]
    for row in rows:
        table.remove(row)
    first_row = copy.deepcopy(rows[0])
    for cell in [node for node in first_row if local_name(node) == "tc"]:
        first_row.remove(cell)
    cell = body_cells[-1]
    first_row.append(cell)
    table.append(first_row)
    for child in cell:
        if local_name(child) == "cellAddr":
            child.set("rowAddr", "0")
            child.set("colAddr", "0")
        elif local_name(child) == "cellSpan":
            child.set("rowSpan", "1")
            child.set("colSpan", "1")
        elif local_name(child) == "cellSz":
            child.set("width", width)
        elif local_name(child) == "subList":
            for cell_paragraph in direct(child, "p"):
                child.remove(cell_paragraph)
            child.append(copy.deepcopy(paragraph_template))
    structural_run = etree.Element(
        "{http://www.hancom.co.kr/hwpml/2011/paragraph}run",
        charPrIDRef=anchor_style,
    )
    structural_run.append(table)
    lines = next(iter(direct(paragraph, "linesegarray")), None)
    position = paragraph.index(lines) if lines is not None else len(paragraph)
    paragraph.insert(position, structural_run)
    section.append(paragraph)
    return table
