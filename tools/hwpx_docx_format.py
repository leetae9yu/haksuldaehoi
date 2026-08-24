from __future__ import annotations

import re

from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree

ALIGNMENTS = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "DISTRIBUTE": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}
FONT_FALLBACK = "Noto Serif CJK KR"


def local_name(element: etree._Element) -> str:
    return etree.QName(element).localname


def first(element: etree._Element | None, name: str) -> etree._Element | None:
    if element is None:
        return None
    return next((node for node in element.iter() if local_name(node) == name), None)


def direct(element: etree._Element, name: str) -> list[etree._Element]:
    return [node for node in element if local_name(node) == name]


def number(element: etree._Element | None, name: str, default: int = 0) -> int:
    if element is None:
        return default
    value = element.get(name)
    return int(value) if value and re.fullmatch(r"-?\d+", value) else default


def hwp_inches(value: int) -> Inches:
    return Inches(value / 7200)


class FormatCatalog:
    def __init__(self, header: etree._Element) -> None:
        hangul_face = next(
            (
                node
                for node in header.iter()
                if local_name(node) == "fontface" and node.get("lang") == "HANGUL"
            ),
            None,
        )
        self.fonts = (
            {
                node.get("id", ""): node.get("face", FONT_FALLBACK)
                for node in hangul_face
                if local_name(node) == "font"
            }
            if hangul_face is not None
            else {}
        )
        self.char_properties = {
            node.get("id", ""): node
            for node in header.iter()
            if local_name(node) == "charPr"
        }
        self.para_properties = {
            node.get("id", ""): node
            for node in header.iter()
            if local_name(node) == "paraPr"
        }
        self.border_fills = {
            node.get("id", ""): node
            for node in header.iter()
            if local_name(node) == "borderFill"
        }

    def apply_run(self, run: Run, style_id: str) -> None:
        style = self.char_properties.get(style_id)
        if style is None:
            return
        font_reference = first(style, "fontRef")
        font_id = font_reference.get("hangul", "") if font_reference is not None else ""
        font_name = self.fonts.get(font_id, FONT_FALLBACK)
        run.font.name = font_name
        run._element.get_or_add_rPr().get_or_add_rFonts().set(
            qn("w:eastAsia"), font_name
        )
        run.font.size = Pt(number(style, "height", 1000) / 100)
        run.bold = first(style, "bold") is not None
        run.italic = first(style, "italic") is not None
        underline = first(style, "underline")
        run.underline = underline is not None and underline.get("type") != "NONE"
        run.font.superscript = first(style, "supscript") is not None
        run.font.subscript = first(style, "subscript") is not None
        spacing = first(style, "spacing")
        spacing_percent = number(spacing, "hangul")
        if spacing_percent:
            spacing_node = OxmlElement("w:spacing")
            spacing_node.set(
                qn("w:val"),
                str(
                    round(
                        number(style, "height", 1000) / 100 * 20 * spacing_percent / 100
                    )
                ),
            )
            run._element.get_or_add_rPr().append(spacing_node)
        color = style.get("textColor", "").removeprefix("#")
        if re.fullmatch(r"[0-9A-Fa-f]{6}", color):
            run.font.color.rgb = RGBColor.from_string(color)

    def apply_paragraph(self, paragraph: Paragraph, style_id: str) -> None:
        style = self.para_properties.get(style_id)
        if style is None:
            return
        alignment = first(style, "align")
        if alignment is not None:
            mapped_alignment = ALIGNMENTS.get(alignment.get("horizontal", ""))
            if mapped_alignment is not None:
                paragraph.alignment = mapped_alignment
        margin = first(style, "margin")
        if margin is not None:
            values = {local_name(node): number(node, "value") for node in margin}
            paragraph.paragraph_format.left_indent = hwp_inches(values.get("left", 0))
            paragraph.paragraph_format.right_indent = hwp_inches(values.get("right", 0))
            paragraph.paragraph_format.space_before = Pt(values.get("prev", 0) / 100)
            paragraph.paragraph_format.space_after = Pt(values.get("next", 0) / 100)
            indent = values.get("intent", values.get("indent", 0))
            paragraph.paragraph_format.first_line_indent = hwp_inches(indent)
        spacing = first(style, "lineSpacing")
        if spacing is not None:
            kind = spacing.get("type", spacing.get("typ", ""))
            value = number(spacing, "value", 160)
            paragraph.paragraph_format.line_spacing = (
                value / 100 if kind == "PERCENT" else Pt(value / 100)
            )

    def apply_cell(self, cell: _Cell, source: etree._Element) -> None:
        cell_size = first(source, "cellSz")
        cell.width = hwp_inches(number(cell_size, "width"))
        properties = cell._tc.get_or_add_tcPr()
        border_fill = self.border_fills.get(source.get("borderFillIDRef", ""))
        if border_fill is not None:
            borders = OxmlElement("w:tcBorders")
            for side in ("left", "right", "top", "bottom"):
                source_border = first(border_fill, f"{side}Border")
                border = OxmlElement(f"w:{side}")
                border_type = (
                    source_border.get("type", "NONE")
                    if source_border is not None
                    else "NONE"
                )
                border.set(qn("w:val"), "nil" if border_type == "NONE" else "single")
                if source_border is not None:
                    color = source_border.get("color", "auto").removeprefix("#")
                    border.set(qn("w:color"), "auto" if color == "none" else color)
                    width = float(source_border.get("width", "0.1 mm").split()[0])
                    border.set(qn("w:sz"), str(max(2, round(width * 72 / 25.4 * 8))))
                borders.append(border)
            properties.append(borders)
            brush = first(border_fill, "winBrush")
            color = brush.get("faceColor", "") if brush is not None else ""
            if re.fullmatch(r"#[0-9A-Fa-f]{6}", color):
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), color[1:])
                properties.append(shading)
        margins = first(source, "cellMargin")
        if margins is not None:
            cell_margins = OxmlElement("w:tcMar")
            for side in ("top", "left", "bottom", "right"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), str(round(number(margins, side) / 5)))
                node.set(qn("w:type"), "dxa")
                cell_margins.append(node)
            properties.append(cell_margins)

    def apply_table(self, table: Table, source: etree._Element) -> None:
        border_fill = self.border_fills.get(source.get("borderFillIDRef", ""))
        if border_fill is None:
            return
        borders = OxmlElement("w:tblBorders")
        for side in ("left", "right", "top", "bottom"):
            source_border = first(border_fill, f"{side}Border")
            if source_border is None or source_border.get("type") == "NONE":
                continue
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(
                qn("w:color"), source_border.get("color", "#000000").removeprefix("#")
            )
            width = float(source_border.get("width", "0.1 mm").split()[0])
            border.set(qn("w:sz"), str(max(2, round(width * 72 / 25.4 * 8))))
            borders.append(border)
        table._tbl.tblPr.append(borders)

    @staticmethod
    def configure_page(document: DocumentObject, section_root: etree._Element) -> None:
        page = first(section_root, "pagePr")
        margin = first(page, "margin")
        if page is None:
            return
        section = document.sections[0]
        section.page_width = hwp_inches(number(page, "width", 59528))
        section.page_height = hwp_inches(number(page, "height", 84186))
        if margin is not None:
            for name in ("left", "right", "top", "bottom", "header", "footer"):
                setattr(section, f"{name}_margin", hwp_inches(number(margin, name)))
        if first(section_root, "pageNum") is not None:
            section.different_first_page_header_footer = True
            paragraph = section.footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("-")
            begin = OxmlElement("w:fldChar")
            begin.set(qn("w:fldCharType"), "begin")
            instruction = OxmlElement("w:instrText")
            instruction.set(qn("xml:space"), "preserve")
            instruction.text = " PAGE "
            end = OxmlElement("w:fldChar")
            end.set(qn("w:fldCharType"), "end")
            run._r.extend([begin, instruction, end])
            paragraph.add_run("-")
