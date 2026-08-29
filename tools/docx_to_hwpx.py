from __future__ import annotations

import argparse
import copy
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

from tools.docx_hwpx_content import (
    direct_text,
    fill_paragraph,
    local_name,
    read_footnotes,
)
from tools.docx_hwpx_layout import (
    aligned_paragraphs,
    deduplicate_tables,
    restore_cover_layout,
    top_paragraphs,
)
from tools.docx_hwpx_styles import CharCatalog
from tools.docx_hwpx_table import append_single_cell_table, fill_table
from tools.hwpx_to_docx import convert_hwpx


def _first_footnote(section: etree._Element) -> etree._Element:
    return next(
        node for node in section.iter() if local_name(node) == "footNote"
    )


def _tables(section: etree._Element) -> list[etree._Element]:
    return [node for node in section.iter() if local_name(node) == "tbl"]


def _write_package(
    reference: Path,
    target: Path,
    section: etree._Element,
) -> None:
    with zipfile.ZipFile(reference) as source, zipfile.ZipFile(target, "w") as output:
        for info in source.infolist():
            payload = source.read(info.filename)
            if info.filename == "Contents/section0.xml":
                payload = etree.tostring(
                    section,
                    encoding="UTF-8",
                    xml_declaration=True,
                )
            elif info.filename == "Preview/PrvText.txt":
                payload = "\n".join(
                    direct_text(paragraph)
                    for paragraph in top_paragraphs(section)
                    if direct_text(paragraph)
                ).encode()
            output.writestr(info, payload)


def convert_docx_to_hwpx(source: Path, reference: Path, target: Path) -> None:
    with tempfile.TemporaryDirectory(prefix="docx-hwpx-") as temporary:
        reference_docx_path = Path(temporary) / "reference.docx"
        convert_hwpx(reference, reference_docx_path)
        reference_docx = Document(str(reference_docx_path))
        source_docx = Document(str(source))
        with zipfile.ZipFile(reference) as archive:
            header = etree.fromstring(archive.read("Contents/header.xml"))
            section = etree.fromstring(archive.read("Contents/section0.xml"))

        templates = top_paragraphs(section)
        aligned = aligned_paragraphs(
            source_docx.paragraphs,
            reference_docx.paragraphs,
            templates,
        )
        notes = read_footnotes(source)
        note_template = _first_footnote(section)
        catalog = CharCatalog(header)
        note_number = [0]
        for paragraph in list(templates):
            section.remove(paragraph)
        for source_paragraph, template in aligned:
            paragraph = copy.deepcopy(template)
            fill_paragraph(
                paragraph,
                source_paragraph,
                catalog,
                notes,
                note_template,
                note_number,
            )
            section.append(paragraph)

        deduplicate_tables(section)
        candidate_tables = _tables(section)
        for table, source_table in zip(
            candidate_tables,
            source_docx.tables,
            strict=False,
        ):
            fill_table(
                table,
                source_table,
                catalog,
                notes,
                note_template,
                note_number,
            )
        restore_cover_layout(section, templates)
        if len(source_docx.tables) > len(candidate_tables):
            table_paragraph = next(
                paragraph
                for paragraph in templates
                if any(local_name(node) == "tbl" for node in paragraph.iter())
            )
            width = max(
                int(
                    next(
                        child
                        for child in table
                        if local_name(child) == "sz"
                    ).get("width", "0")
                )
                for table in candidate_tables
            )
            body_paragraph = next(
                paragraph
                for paragraph in templates
                if paragraph.get("paraPrIDRef") == "11"
                and direct_text(paragraph).strip()
                and not any(
                    local_name(node) == "footNote"
                    for node in paragraph.iter()
                )
            )
            added_table = append_single_cell_table(
                section,
                table_paragraph,
                candidate_tables[-1],
                body_paragraph,
                str(max(int(table.get("id", "0")) for table in candidate_tables) + 1),
                str(width),
            )
            fill_table(
                added_table,
                source_docx.tables[-1],
                catalog,
                notes,
                note_template,
                note_number,
            )
        target.parent.mkdir(parents=True, exist_ok=True)
        _write_package(reference, target, section)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert DOCX to formatted HWPX.")
    parser.add_argument("source", type=Path)
    parser.add_argument("reference", type=Path)
    parser.add_argument("target", type=Path)
    arguments = parser.parse_args()
    convert_docx_to_hwpx(arguments.source, arguments.reference, arguments.target)


if __name__ == "__main__":
    main()
