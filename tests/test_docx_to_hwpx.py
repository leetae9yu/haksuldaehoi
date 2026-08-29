from __future__ import annotations

import copy
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

from tools.docx_hwpx_content import fill_paragraph, local_name
from tools.docx_hwpx_styles import CharCatalog
from tools.docx_to_hwpx import convert_docx_to_hwpx
from tools.hwpx_format_map import compare_format_maps
from tools.hwpx_to_docx import convert_hwpx
from tools.validate_docx_hwpx import validate_documents as validate_reverse
from tools.validate_hwpx_docx import validate_documents


def test_manual_docx_break_does_not_stretch_justified_hwpx_line() -> None:
    reference = Path("본논문_1차.hwpx")
    with zipfile.ZipFile(reference) as archive:
        header = etree.fromstring(archive.read("Contents/header.xml"))
        section = etree.fromstring(archive.read("Contents/section0.xml"))
    target = copy.deepcopy(
        next(
            node
            for node in section.iter()
            if local_name(node) == "p"
            and "".join(str(text) for text in node.itertext()).strip()
        )
    )
    note_template = next(
        node for node in section.iter() if local_name(node) == "footNote"
    )
    document = Document()
    source = document.add_paragraph()
    run = source.add_run("앞")
    run.add_break()
    run.add_text("뒤")

    fill_paragraph(
        target,
        source,
        CharCatalog(header),
        {},
        note_template,
        [0],
    )

    assert all(
        local_name(node) != "lineBreak"
        for node in target.iter()
    )


def test_hwpx_docx_hwpx_roundtrip_preserves_content_and_format(
    tmp_path: Path,
) -> None:
    reference = Path("본논문_1차.hwpx")
    source_docx = tmp_path / "source.docx"
    candidate = tmp_path / "candidate.hwpx"
    roundtrip_docx = tmp_path / "roundtrip.docx"

    convert_hwpx(reference, source_docx)
    convert_docx_to_hwpx(source_docx, reference, candidate)
    convert_hwpx(candidate, roundtrip_docx)

    assert compare_format_maps(reference, candidate) == []
    assert validate_reverse(source_docx, reference, candidate)["valid"] is True
    assert validate_documents(candidate, roundtrip_docx)["valid"] is True
